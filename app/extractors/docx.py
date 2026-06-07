from __future__ import annotations
from io import BytesIO
from pathlib import Path
from typing import Union
from app.extractors.base import Extractor


class DocxExtractor(Extractor):
    def extract(self, source: Union[str, Path]) -> str:
        return self.extract_bytes(Path(source).read_bytes())

    def extract_bytes(self, data: bytes) -> str:
        from docx import Document

        doc = Document(BytesIO(data))
        parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts).strip()
